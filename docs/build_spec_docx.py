#!/usr/bin/env python3
"""Generate the DemoTrack project specification as a styled .docx."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = "5B3DF5"
PURPLE_DK = "4A2FD0"
INK = "1C2230"
SLATE = "2A2F45"
GREY = "6B7280"
GOLD = "B8860B"
GREEN = "1A7F37"
BLUE = "2B7FFF"
LAV_BG = "F1EEFF"
NOTE_BG = "EAF3FF"
WARN_BG = "FFF7E0"
ZEBRA = "F6F6FB"
HDR = PURPLE
BORDER = "D6D9E0"

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
pf = normal.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.12

sec = doc.sections[0]
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)


def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    el.append(e)
    return e


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    _set(pPr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})


def left_bar(paragraph, color, size="24"):
    pPr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "right"):
        _set(bdr, f"w:{side}", **{"w:val": "single", "w:sz": "2", "w:space": "0", "w:color": fill_bg})
    _set(bdr, "w:left", **{"w:val": "single", "w:sz": size, "w:space": "8", "w:color": color})
    pPr.append(bdr)


def bottom_rule(paragraph, color, size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    _set(bdr, "w:bottom", **{"w:val": "single", "w:sz": size, "w:space": "4", "w:color": color})
    pPr.append(bdr)


# inline markdown: **bold**, *italic*, `code`
_TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def render_inline(paragraph, text, base_color=INK, base_size=None, base_bold=False):
    for part in _TOKEN.split(text):
        if not part:
            continue
        bold, italic, mono = base_bold, False, False
        t = part
        if part.startswith("**") and part.endswith("**"):
            bold, t = True, part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            italic, t = True, part[1:-1]
        elif part.startswith("`") and part.endswith("`"):
            mono, t = True, part[1:-1]
        r = paragraph.add_run(t)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = RGBColor.from_string(base_color)
        if base_size:
            r.font.size = Pt(base_size)
        if mono:
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
    return paragraph


def body(text, space_after=6, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    render_inline(p, text, base_color=color)
    return p


def h_section(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    bottom_rule(p, PURPLE, "16")
    r = p.add_run(f"{num}.  {text}" if num else text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    return p


def h4(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    rPr = r._element.get_or_add_rPr()
    _set(rPr, "w:spacing", **{"w:val": "30"})
    return p


fill_bg = "FFFFFF"


def callout(lead, text, bg, bar):
    global fill_bg
    fill_bg = bg
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(4)
    shade(p, bg)
    left_bar(p, bar)
    if lead:
        r = p.add_run(lead + "  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
    render_inline(p, text)
    return p


def bullets(items, style_color=INK):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        render_inline(p, it, base_color=style_color)


def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        render_inline(p, it)


def _cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    _set(tcPr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        _set(m, f"w:{side}", **{"w:w": str(val), "w:type": "dxa"})
    tcPr.append(m)


def _table_borders(tbl, color=BORDER):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set(borders, f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": color})
    tblPr.append(borders)


def make_table(headers, rows, widths=None, header_fill=HDR, zebra=True, font=9.5):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _table_borders(tbl)
    # header
    hcells = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        _cell_shade(hcells[i], header_fill)
        _cell_margins(hcells[i])
        p = hcells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
    # body
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            _cell_margins(cells[ci])
            if zebra and ri % 2 == 1:
                _cell_shade(cells[ci], ZEBRA)
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            render_inline(p, val, base_size=font)
    if widths:
        for ci, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def page_break():
    doc.add_page_break()


def spacer(pts=4):
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)


# ============================ COVER ============================
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(0)
r = t.add_run("DemoTrack")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = RGBColor.from_string(PURPLE)

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("Your silent manager.")
r.italic = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string(PURPLE_DK)

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(10)
bottom_rule(meta, BORDER, "6")
r = meta.add_run("Project Specification  ·  v3.1 (final)  ·  House & Tech House  ·  Independent Artist Toolkit  ·  Tunis, Tunisia  ·  2026")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor.from_string(GREY)

body("A career-management app for a solo house & tech-house producer with no team. "
     "It strips the friction out of getting your music to the right labels — turning demo "
     "submission from a thing you keep meaning to do into a **weekly habit that compounds.**",
     space_after=10)

# stats row
stats = [("13", "Modules"), ("12", "Data Tables"), ("13", "Build Phases"), ("17", "Labels Seeded"), ("3", "AI Features")]
st = doc.add_table(rows=2, cols=len(stats))
st.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (n, l) in enumerate(stats):
    c = st.cell(0, i)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(n)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    c2 = st.cell(1, i)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(l.upper())
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor.from_string(GREY)
spacer(8)

callout("How to read this document.",
        "It explains, in plain terms, *what* DemoTrack is, *how* it works, the data behind it, "
        "the technology, and the order it gets built in. Every design choice traces back to one idea: "
        "a solo artist's bottleneck is **effort and consistency**, not talent — so the whole app is "
        "engineered to remove friction and manufacture momentum.", NOTE_BG, BLUE)

page_break()

# ============================ 1 ============================
h_section("1", "What it is & why it exists")
body("DemoTrack is a personal command center for breaking through as an unknown producer. It does "
     "the job a manager would do for a signed artist — holding the accountability, the memory, and the "
     "workflow — because you don't have a manager.")
h3("The problem")
body("Getting signed is mostly an outreach game: finding the right labels, sending demos the right way, "
     "following up at the right moment, and learning from every response. The work isn't intellectually "
     "hard — it's tedious, repetitive, and very easy to skip. The typical pattern: a producer sends a "
     "burst of demos once, hears nothing for a few weeks, gets discouraged, and the pipeline goes cold. "
     "No system, no memory, no momentum.")
callout("The core insight that shaped everything.",
        "The bottleneck isn't talent and usually isn't even contacts — it's effort and consistency. If "
        "sending a polished, tailored demo takes 3 minutes instead of 30, and the app proactively tells "
        "you who to follow up with today, you will actually do the work.", WARN_BG, GOLD)
callout("One-line pitch:",
        "“Find good labels, send great demos in minutes, never lose track of a follow-up, and learn "
        "from every reply — automatically.”", LAV_BG, PURPLE)
h3("The quality-over-quantity rule")
body("Only labels that can genuinely change your trajectory belong in the system. A label clears the bar "
     "only if it meets **at least 3 of 5** tests: Beatport chart presence · booking impact · streaming "
     "credibility · scene respect · accessible to unknown artists. This filter is applied to every "
     "seeded label.")

# ============================ 2 ============================
h_section("2", "How it works — the core loop")
body("Everything serves one repeating cycle. You run this loop for every track you want to place. The app "
     "makes each step faster and remembers the state, so nothing falls through the cracks.")
make_table(
    ["1 Discover", "2 Prepare", "3 Generate", "4 Send", "5 Track", "6 Follow-up", "7 Learn"],
    [["find labels", "track + kit", "preset + hook", "< 3 min", "opens / seen", "right time", "log + patterns"]],
    zebra=False, font=8.5)
body("**The loop is the product.** Modules are just the tools that make each step effortless. After step 7 "
     "you return to step 1 with the next track.")

# ============================ 3 ============================
h_section("3", "The structure — four layers")
make_table(["Layer", "What lives there", "Purpose"],
    [["🎵 Music", "Track Vault, Feedback Log", "Your catalogue and its memory — what you're shopping and how it lands."],
     ["📨 Outreach (the heart)", "Contacts CRM, Send Demo, Email Presets + AI Hook, Follow-up, A&R Intel, Label Discovery, (optional) Link Tracking", "Where demos go out, get tracked, and get followed up."],
     ["👤 You", "Artist Press Kit, Work Sessions, Goals", "The solo-operator backbone — your EPK, accountability, and targets."],
     ["📊 Dashboard", "Home screen + Notifications digest", "Pulls signals from every module and answers “what should I do right now?”"]],
    widths=[1.5, 3.0, 2.6])
callout("Note on prioritization:",
        "there is **no weighted “contact score.”** A computed score has a cold-start problem "
        "(almost no data early) and is false precision you'd never trust. Instead, each label simply shows "
        "a **“times contacted” count** (derived from your send history) and an **access path** "
        "(how reachable it realistically is). Untried, cold-demo-friendly labels surface first. Nothing to "
        "compute, nothing that breaks on day one.", NOTE_BG, BLUE)

page_break()

# ============================ 4 ============================
h_section("4", "The 13 modules in detail")
h4("Music Layer")
h3("01 · Track Vault")
body("Your library of finished and in-progress tracks, each with a status: idea → demo-ready → submitted → signed.")
bullets([
    "**Key fields:** title, genre tags, BPM, key, status, a private listen link + link type, notes, readiness checklist, exclusive-hold (which label is considering it, until when).",
    "**Audio is never hosted by the app** — you store a link (unlisted SoundCloud recommended, which also gives you free play counts). Keeps infrastructure light.",
    "**Why it matters:** the single source of truth for “what am I shopping right now.” Every submission and piece of feedback links back to a track here.",
])
h3("02 · Feedback Log")
body("Captures every response a track gets — yes, no, “not for us,” constructive notes, or silence — tied to the track and the contact.")
bullets([
    "**Silence is data:** if a demo goes unanswered past the overdue window, the app **auto-logs a “no-response” record** — you type nothing.",
    "**Real replies** are logged in one tap from the follow-up queue (a response-type dropdown + optional pasted text).",
    "**Why it matters:** over time the AI reads across all your feedback and surfaces patterns, telling you what to make next and who to send it to.",
])
h4("Outreach Layer")
h3("03 · Contacts CRM")
body("Every label and industry contact in one place, across 7 categories: Labels, DJs, A&Rs, Curators, Blogs, Promoters, Radio.")
bullets([
    "**Key fields:** name, category, submission method, email/portal/DM link, relationship stage (cold → engaged → responded → relationship), access path, last contacted, full history, and a back-reference to the Discovery master it came from.",
    "**Why it matters:** knows who you've already contacted so you never double-send or forget someone. The backbone the entire outreach layer reads from.",
])
h3("04 · Send Demo")
body("The submission flow and the most important screen in the app. Pick a track, pick a contact, choose a preset, the app auto-fills the email, you send — under 3 minutes.")
bullets([
    "**Adapts to the method:** Email opens prefilled in your own inbox; Form opens the portal; DM opens the profile. You confirm so every send is logged with a status.",
    "**Pre-send checklist:** the target's submission requirements render as a checklist, and the app **warns if the track is on exclusive hold elsewhere.**",
    "**Why it matters:** this is the friction-killer the whole product is built around. If this flow is fast and pleasant, you'll submit weekly.",
])
h3("05 · Email Presets + AI Hook")
body("The email is **not** AI-written end-to-end. You keep your own preset templates; the app auto-fills "
     "everything mechanical, and AI suggests only the one personalized sentence that earns a reply. Detailed in Section 7.")
h3("06 · Follow-up Workflows")
body("Automatically queues the right follow-up at the right time.")
bullets([
    "**Timing:** 7-day silence nudge · 14-day overdue flag.",
    "**One-tap reply status:** because the app sends from your own inbox and can't read replies, the queue shows each open send with two taps — **“Got a reply”** or **“Still silent.”** No nagging someone who already answered.",
    "**Why it matters:** following up is where demos actually get signed — and it's the single easiest thing to forget.",
])
h3("07 · A&R Intel")
body("Research notes on each label — who runs it, what they sign, recent releases, submission preferences, "
     "and any personal angle. One record per contact. This is also what feeds the AI's personalized hook. "
     "Context turns a generic blast into informed outreach.")
h3("08 · Label Discovery")
body("A browsable, filterable database of career-changing labels — separate from your personal CRM — so you always know your next target.")
bullets([
    "**Filters:** genre, BPM range, submission method, tier, access path, last-verified date.",
    "**“Add to my CRM”** copies a label into your contacts with a link back to the master; if the master is re-verified, your copy is flagged “source updated — review.”",
    "**Why it matters:** removes “who do I send to next?” as a blocker.",
])
h3("09 · Demo Link Tracking  (optional · Phase 11)")
body("Wraps demo links so you can see when a label opens your demo. **Scoped to Form/DM links only — never "
     "email** (email open-tracking is unreliable and hurts deliverability). Lower priority because unlisted "
     "SoundCloud already gives you a real listen signal for free.")
h4("You Layer")
h3("10 · Artist Press Kit")
body("A public EPK page at press.demotrack.app/{your-slug} with bio, links, releases, and stats — plus AI bio generation in 3 tones.")
bullets([
    "**Stats are manual** with an “updated {date}” stamp (no fragile third-party integrations).",
    "**Served safely:** the public page is rendered through a narrow slug endpoint that returns only whitelisted fields — your private CRM and tracks can never leak.",
    "**Auto-attach toggle** adds the kit link to every demo by default.",
])
h3("11 · Work Sessions")
body("Logs studio time — but kept to **hours + a one-tap mood** (fire / good / ok / low); “demos sent "
     "that day” is auto-filled. The gentle accountability a manager would otherwise provide, with almost no friction.")
h3("12 · Goals")
body("Targets with live progress — e.g. “send 4 demos/month.” **Progress is derived automatically** "
     "from your real activity (you only set the target). Turns vague ambition into measurable habits.")
h3("★ Dashboard")
body("The home screen and connective tissue. Shows: demos to send, follow-ups due today, recently seen demos, "
     "goal progress, recent feedback, your work-session streak, and a **funnel** (sent → opened → replied → "
     "considering → signed) so you can see your real conversion and stay realistic.")

page_break()

# ============================ 5 ============================
h_section("5", "The three AI features")
body("Claude is used in three specific, high-leverage places — to remove the exact tasks that block a solo producer, not as a gimmick.")
make_table(["Feature", "What it does", "Guardrails"],
    [["1. Email hook suggestion", "Writes only the single tailored sentence (“why this track fits this label”) that slots into your preset.", "≤30 words; never fabricates a release/stat/angle; required; warns on duplicates."],
     ["2. Artist bio generation", "Writes the press-kit bio in 3 selectable tones from your profile facts.", "Facts only — no invented numbers or accolades."],
     ["3. Feedback pattern detection", "Reads across your whole feedback log and surfaces which sounds/labels/approaches get the best response.", "Needs a minimum number of responses before claiming a pattern; cites its evidence; counts silence."]],
    widths=[1.7, 3.0, 2.4])
callout("Why AI here specifically.",
        "Each of these is a task that is boring, repeatable, and easy to put off — exactly the kind of work "
        "that stalls a solo artist. With email trimmed to just the hook, AI usage and cost stay tiny.", WARN_BG, GOLD)

# ============================ 6 ============================
h_section("6", "How outreach works")
h3("Three submission methods")
make_table(["Method", "Badge", "Flow"],
    [["Email", "🟢 Green", "Pick preset → app auto-fills merge fields → AI suggests the hook (you tweak) → **opens your own Gmail / mail app, prefilled** → you send → confirm logged. No in-app relay = no spam risk."],
     ["Form only", "🟡 Amber", "App opens the portal/platform in a new tab → you submit manually → confirm in the app so it's logged."],
     ["DM only", "🔵 Blue", "App opens the profile → you send the DM → confirm so it's tracked."]],
    widths=[1.1, 1.0, 5.0])
h3("Prioritization — two honest signals, zero cold-start")
bullets([
    "**Access path** — who's even worth a cold send. Discovery shows cold-demo-friendly first; relationship-only elites go to a separate “long game” list.",
    "**Times contacted** — a simple count per label. 0 = your next target; a high count with no reply = change the track or move on.",
])
body("**Default sort:** cold-demo-friendly + untried first, then by tier. Stale labels get a freshness badge.")
h3("Follow-up timing & reply visibility")
bullets([
    "**7 days silence** → gentle first-nudge queued.",
    "**14 days** → flagged overdue; a “no-response” feedback record auto-logged.",
    "The queue is driven by a send **status** you advance with one tap — never by reading your inbox.",
    "All of it is pushed to you in a **scheduled email digest**, not just parked in-app.",
])

page_break()

# ============================ 7 ============================
h_section("7", "Email presets + AI hook — in detail")
body("This is the most-used flow, so it's worth understanding fully. A preset is a **skeleton with merge "
     "fields** the app fills automatically, plus **exactly one slot you personalize.**")
# mono code block as shaded paragraph
code_lines = [
    "Subject: {genre} demo — “{track_title}” for {label}",
    "",
    "Hi {first_name},",
    "",
    "I'm {artist_name}, a house & tech house producer from Tunis. {hook}",
    "",
    "“{track_title}” — {bpm} BPM, {key}. Private listen: {listen_link}",
    "More on me: {press_kit_link}",
    "",
    "Would love your thoughts. Thanks for listening.",
    "{artist_name}",
]
cp = doc.add_paragraph()
cp.paragraph_format.left_indent = Pt(8)
cp.paragraph_format.space_before = Pt(4)
cp.paragraph_format.space_after = Pt(8)
shade(cp, "1C2230")
for i, line in enumerate(code_lines):
    run = cp.add_run(line + ("\n" if i < len(code_lines) - 1 else ""))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("E8E6F5")
body("Everything in {...} auto-fills from the track, the contact, and your profile — **except {hook}.** AI "
     "proposes that one sentence (e.g. *“Your recent Cloonee EP on Hellbent is exactly the lane this track "
     "sits in”*); you confirm or tweak it, and the filled email opens in your own inbox to send. A send "
     "becomes: **pick track → pick label → confirm one sentence → send.**")
body("**Anti-mass-blast guardrail:** the hook is required and the app warns if it's empty or identical to a "
     "recent send. Suggested starter presets: *cold email*, *DM* (shorter), *warm follow-up*, *form-portal note*.")

# ============================ 8 ============================
h_section("8", "Design principles (apply to every module)")
body("The whole app exists to remove friction, so manual data entry is the enemy. These rules are non-negotiable:")
numbered([
    "**Derive, don't ask.** If a number can be computed from existing data, never make the user type it (times contacted, goal progress, demos-sent, the funnel are all computed).",
    "**One-tap over typing.** Reply status, mood, send confirmation are single taps. Free text is reserved for the only three things a human must write: the **hook**, **A&R intel notes**, and **feedback text.**",
    "**Silence is data.** An unanswered demo past the overdue window is logged automatically.",
    "**Structural privacy.** Public data (the press kit) is exposed through a narrow whitelisted endpoint, not by trusting a permission rule to be perfect.",
])

page_break()

# ============================ 9 ============================
h_section("9", "The data model — 12 tables")
body("All data is stored in Supabase (PostgreSQL) with Row-Level Security, so every row is private to your account.")
make_table(["Table", "Holds"],
    [["contacts", "Your personal CRM — every label/DJ/A&R/curator. Carries relationship stage, access path, and a back-reference to the Discovery master."],
     ["tracks", "Your catalogue and each track's status, listen link + type, readiness checklist, and any exclusive hold."],
     ["submissions", "Every demo sent (links a track to a contact); method, follow-up dates, and a status (sent → opened → replied → considering → signed → passed)."],
     ["feedback", "Every response, tied to a track and contact. Silence is auto-logged as a no-response record."],
     ["templates", "Your preset email & follow-up skeletons with merge fields and one hook slot."],
     ["ar_intel", "Research notes — one record per contact."],
     ["work_sessions", "Studio time logs (hours + one-tap mood; demos-sent auto-filled)."],
     ["goals", "Your targets. Progress is a derived view, not a stored number."],
     ["link_events", "Each click on a tracked Form/DM link (Phase 11, optional)."],
     ["press_kit", "Your public EPK profile — one row per user; served via a whitelisted slug endpoint."],
     ["labels", "The curated discovery library (separate from your CRM): access path, genre tags, last-verified, sources, sub-label links, submission requirements."],
     ["notifications", "Queued/sent reminders & digests, fired by a scheduler."]],
    widths=[1.4, 5.7])
body("**How they connect:** your account owns everything. A contact has many submissions; each submission can "
     "have feedback (and, later, link events). A track has many submissions and feedback, and at most one active "
     "hold. The press kit is one-to-one with you. The labels library is read-only and seeds Discovery; labels you "
     "pursue become contacts. “Times contacted” is just a count of submissions — there is no scoring table.")

page_break()

# ============================ 10 ============================
h_section("10", "The tech stack")
body("A lean, modern stack chosen so one person can build and run it without a team.")
make_table(["Technology", "Role & why"],
    [["React 19 + Vite", "The interface + a fast dev/build environment. Mobile-first, installable PWA (the habit happens on a phone)."],
     ["Tailwind CSS 4", "Styling system — consistent, quick, no CSS sprawl."],
     ["Vercel", "Hosting/deploy for the app and the public press-kit pages; handles the subdomains."],
     ["Supabase (PostgreSQL)", "Database + magic-link auth + Row-Level Security in one. 12 tables."],
     ["Supabase Storage", "One bucket, for the press-kit photo only. Audio is never hosted."],
     ["Supabase pg_cron + Scheduled Edge Function", "Fires the daily/weekly digest and the 7/14-day follow-up triggers."],
     ["Anthropic API (Claude)", "Powers the 3 AI features. With email trimmed to the hook, usage and cost are tiny."],
     ["mailto: / Gmail compose deep link", "Demo email opens prefilled in your own inbox — no relay, no send engine, no deliverability risk."],
     ["Edge Function (Deno)", "The optional Phase-11 link-tracking redirect (Form/DM only) and the slug endpoint serving the public press kit."]],
    widths=[2.2, 4.9])
body("**Domains:** demotrack.app (the app) · press.demotrack.app (public press kit) · track.demotrack.app "
     "(Phase-11 redirects, Form/DM links only).")

# ============================ 11 ============================
h_section("11", "The seed database — 17 curated labels")
body("Loaded in Phase 3. Every one clears the quality bar. **Tier** reflects prestige; **access path** reflects "
     "how realistically reachable it is for an unknown — they are different axes.")
seed = [
    ["Drumcode", "ELITE", "relationship-only", "Adam Beyer. Scene-defining — a long game, not a cold target."],
    ["Experts Only", "ELITE", "needs-warm-intro", "John Summit. Peak of the scene."],
    ["Black Book", "ELITE", "needs-warm-intro", "Chris Lake. A Beatport #1 machine."],
    ["Catch & Release", "ELITE", "needs-warm-intro", "Fisher. Massive reach and chart dominance."],
    ["Repopulate Mars", "A", "cold-demo-friendly", "Lee Foss. Where Mau P broke through. Open to demos."],
    ["Dirtybird", "A", "cold-demo-friendly", "Claude VonStroke. Fun, quirky lane; launches careers."],
    ["Hellbent", "A", "cold-demo-friendly", "Cloonee. Built for unknowns; charts constantly."],
    ["Three Six Zero", "A", "cold-demo-friendly", "Dom Dolla. Open submissions; growing fast."],
    ["Hot Creations", "A", "needs-warm-intro", "Jamie Jones. Deep scene credibility."],
    ["Solid Grooves", "A", "cold-demo-friendly", "Michael Bibi. UK tech house institution."],
    ["Sola", "A", "cold-demo-friendly", "Solardo. Consistent charting; email demos welcome."],
    ["Toolroom", "B", "cold-demo-friendly", "Mark Knight. Solid distribution; long track record."],
    ["Defected", "B", "cold-demo-friendly", "Huge catalogue; one send covers ~11 sub-labels."],
    ["Knee Deep In Sound", "B", "cold-demo-friendly", "Hot Since 82. A less-competitive melodic lane."],
    ["Suara", "B", "cold-demo-friendly", "Coyu. Darker tech house niche."],
    ["Club Sweat", "B", "cold-demo-friendly", "Australian reach; good for emerging artists."],
    ["Space Yacht", "B", "open-window-only", "Reviews demos live on Twitch — the most transparent route in the scene."],
]
rows = [[f"**{n}**  ·  {tier}", ap, why] for (n, tier, ap, why) in seed]
make_table(["Label & tier", "Access path", "Why it belongs"], rows, widths=[1.9, 1.6, 3.6])
p = doc.add_paragraph()
r = p.add_run("Access-path values are starting suggestions to verify during research — demo windows open and close.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string(GREY)

page_break()

# ============================ 12 ============================
h_section("12", "Building the contact database — a parallel track")
body("The database is the **fuel**; the app is the **engine.** Finding, verifying, and qualifying contacts "
     "that can actually change your career is a research workstream in its own right. It runs **alongside** the "
     "build, not after, and is never fully “done.”")
body("**Research loop per contact:** Source → Verify (2 independent sources) → Qualify (the ≥3/5 bar) → Tier + "
     "access path + priority → Record (rules, links, sources, A&R note) → Maintain (stamp last-verified; "
     "Discovery flags stale entries).")
body("**What already exists:** ≈170 verified labels and ≈98 non-label promotional contacts (radio, blogs, "
     "YouTube/Spotify curators, repost networks, podcasts, editorial, PR), with the 17-label in-app seed selected "
     "from them. **Ongoing:** migrate into the labels table, keep verifications fresh, capture sub-label "
     "efficiency multipliers (1 Defected ≈ 11 sub-labels; 1 Repopulate Mars ≈ 3), and add new qualified contacts.")

# ============================ 13 ============================
h_section("13", "The build plan — phases")
body("Built in this order so the app is usable early and grows in value. Each phase is independently deployable. "
     "**By the end of Phase 4 you can send and record a real demo.**")
h4("Parallel track — runs throughout")
body("**Phase 0 · Contact Database Research** (ongoing) — run the research loop; migrate verified contacts into "
     "the labels table; keep verifications fresh.")
h4("MVP — reach your first real send")
make_table(["Phase", "What you build", "Done when…"],
    [["1 · Foundation", "Supabase project, 12-table schema, magic-link auth, RLS, React+Vite+Tailwind PWA shell, Vercel deploy.", "You can log in, your data is secured and deployed."],
     ["2 · Track Vault", "Add/edit/list tracks, status, external listen link + type, notes, readiness checklist, hold field.", "You can store your tracks and see their status."],
     ["3 · Contacts CRM", "Contact profiles across 7 categories, method + links, relationship stage, access path, seed the 17 labels, history view.", "Your label list lives in the app with submission details."],
     ["4 · Send Demo", "The submission flow — pick track + contact, choose preset, all 3 methods, pre-send checklist + hold warning, record to history.", "**You can send a demo and it's recorded. ← first usable milestone**"],
     ["5 · Email Presets + AI Hook", "Preset CRUD with merge fields, auto-fill, Claude hook suggestion with guardrails, mailto/Gmail prefill.", "A tailored, personal demo email in ~1 minute, from your own inbox."],
     ["6 · Follow-up + Digest", "Preset follow-ups, 7/14-day rules via pg_cron, one-tap reply status, auto-logged silence, overdue queue, scheduled email digest.", "The app tells you who to follow up — and reminds you when it's closed."]],
    widths=[1.6, 3.6, 1.9], font=9)
h4("Intelligence — make every send smarter")
make_table(["Phase", "What you build"],
    [["7 · Feedback Log", "One-tap reply logging, response types, auto no-response records, links back to track & contact."],
     ["8 · A&R Intel & Prioritization", "Intel records per contact, “times contacted” on every label, untried + cold-demo-friendly sort. (No weighted score.)"],
     ["9 · Dashboard + Funnel", "Aggregate signals, the conversion funnel, response-rate by genre/tier, streak, CSV export."]],
    widths=[2.3, 4.8], font=9)
h4("Advanced — the unfair advantages")
make_table(["Phase", "What you build"],
    [["10 · Label Discovery", "Browse/filter the labels master, “Add to my CRM” promotion, “have I submitted?” indicator, freshness warnings, sub-label reach hints."],
     ["11 · Demo Link Tracking (optional)", "Edge Function redirect, unique hash, link events — Form/DM links only."],
     ["12 · Work Sessions & Goals", "Sessions (hours + one-tap mood, demos auto-filled); goals with derived progress."],
     ["13 · Artist Press Kit", "Profile fields, manual stats + updated stamp, public page via slug endpoint, AI bio (3 tones), auto-attach toggle."]],
    widths=[2.3, 4.8], font=9)

page_break()

# ============================ 14 ============================
h_section("14", "What success looks like")
body("The app only matters if it changes your behavior. These are the real measures — not vanity metrics:")
bullets([
    "You submit demos **weekly**, not “whenever I get around to it.”",
    "You **actually follow up** on demos you've sent (because the digest reaches you).",
    "You always know which labels you **haven't tried** yet.",
    "Sending a demo takes **under 3 minutes.**",
    "You get a response from at least one label **within 3 months** of consistent use.",
    "The longer-term target: **one signing within 6 months** of using it consistently.",
])
callout("A note on expectations.",
        "Cold-demo reply rates in this scene are realistically in the 5–15% range, and “we'll consider it” "
        "→ an actual signing is a fraction of that. That is precisely why every send is engineered to cost 3 "
        "minutes: the math only closes when friction is near zero. The funnel view is built to read a quiet month "
        "as “keep going,” not “this is broken.”", WARN_BG, GOLD)

# ============================ 15 ============================
h_section("15", "The unfair advantages")
bullets([
    "**AI woven into the loop** — the hook, bios, and feedback patterns remove the exact tasks that stall solo artists, at near-zero cost.",
    "**A curated, not exhaustive, label database** — only labels that can change your career, kept fresh, with honest access paths and submission rules built in.",
    "**Friction engineered out** — own-inbox sending, derived data, one-tap everything, and a reminder that reaches you. The app does the manager's remembering so you only have to do the music.",
])
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(18)
r = foot.add_run("DemoTrack · Project Specification v3.1 · Built for one producer. One purpose. Ship your music. · Tunis, 2026")
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor.from_string("9AA0AB")

out = "/home/user/DemoTrack/docs/DemoTrack_Specification.docx"
doc.save(out)
print("Saved", out)
