#!/usr/bin/env python3
"""Generate the DemoTrack build plan (phases & detailed tasks) as a styled .docx."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = "5B3DF5"; PURPLE_DK = "4A2FD0"; INK = "1C2230"; SLATE = "2A2F45"
GREY = "6B7280"; GOLD = "B8860B"; GREEN = "1A7F37"; BLUE = "2B7FFF"
LAV_BG = "F1EEFF"; NOTE_BG = "EAF3FF"; WARN_BG = "FFF7E0"; DONE_BG = "EAF7EE"
ZEBRA = "F6F6FB"; HDR = PURPLE; BORDER = "D6D9E0"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12
sec = doc.sections[0]
sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

fill_bg = "FFFFFF"

def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    el.append(e); return e

def shade(p, fill):
    _set(p._p.get_or_add_pPr(), "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})

def left_bar(p, color, size="24"):
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "right"):
        _set(bdr, f"w:{side}", **{"w:val": "single", "w:sz": "2", "w:space": "0", "w:color": fill_bg})
    _set(bdr, "w:left", **{"w:val": "single", "w:sz": size, "w:space": "8", "w:color": color})
    pPr.append(bdr)

def bottom_rule(p, color, size="12"):
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr")
    _set(bdr, "w:bottom", **{"w:val": "single", "w:sz": size, "w:space": "4", "w:color": color})
    pPr.append(bdr)

_TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
def render_inline(p, text, base_color=INK, base_size=None, base_bold=False):
    for part in _TOKEN.split(text):
        if not part:
            continue
        bold, italic, mono, t = base_bold, False, False, part
        if part.startswith("**") and part.endswith("**"): bold, t = True, part[2:-2]
        elif part.startswith("*") and part.endswith("*"): italic, t = True, part[1:-1]
        elif part.startswith("`") and part.endswith("`"): mono, t = True, part[1:-1]
        r = p.add_run(t); r.bold = bold; r.italic = italic
        r.font.color.rgb = RGBColor.from_string(base_color)
        if base_size: r.font.size = Pt(base_size)
        if mono: r.font.name = "Consolas"; r.font.size = Pt(9.5)
    return p

def body(text, space_after=6, color=INK):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    render_inline(p, text, base_color=color); return p

def h_section(text, num=""):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    bottom_rule(p, PURPLE, "16")
    r = p.add_run(f"{num}.  {text}" if num else text)
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(PURPLE); return p

def phase_head(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    bottom_rule(p, PURPLE, "12")
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(PURPLE); return p

def h4(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(PURPLE)
    _set(r._element.get_or_add_rPr(), "w:spacing", **{"w:val": "30"}); return p

def callout(lead, text, bg, bar):
    global fill_bg; fill_bg = bg
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(10); p.paragraph_format.right_indent = Pt(4)
    shade(p, bg); left_bar(p, bar)
    if lead:
        r = p.add_run(lead + "  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(INK)
    render_inline(p, text); return p

def tasks(items):
    for it in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(14); p.paragraph_format.first_line_indent = Pt(-14)
        box = p.add_run("☐  "); box.font.size = Pt(10.5); box.font.color.rgb = RGBColor.from_string(PURPLE)
        render_inline(p, it)

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
        render_inline(p, it)

def _cell_shade(cell, fill):
    _set(cell._tc.get_or_add_tcPr(), "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})

def _cell_margins(cell, t=40, b=40, l=80, r=80):
    m = OxmlElement("w:tcMar")
    for side, val in (("top", t), ("bottom", b), ("start", l), ("end", r)):
        _set(m, f"w:{side}", **{"w:w": str(val), "w:type": "dxa"})
    cell._tc.get_or_add_tcPr().append(m)

def _table_borders(tbl, color=BORDER):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set(borders, f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": color})
    tbl._tbl.tblPr.append(borders)

def make_table(headers, rows, widths=None, zebra=True, font=9.5):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; _table_borders(tbl)
    hc = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _cell_shade(hc[i], HDR); _cell_margins(hc[i])
        p = hc[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(font); r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            _cell_margins(cells[ci])
            if zebra and ri % 2 == 1: _cell_shade(cells[ci], ZEBRA)
            p = cells[ci].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            render_inline(p, val, base_size=font)
    if widths:
        for ci, w in enumerate(widths):
            for row in tbl.rows: row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return tbl

def meta_box(goal, depends, milestone):
    make_table(["", ""], [["Goal", goal], ["Depends on", depends], ["Done when", f"**{milestone}**"]],
               widths=[1.3, 5.8], zebra=False, font=9.5)

def page_break(): doc.add_page_break()

# ============================ COVER ============================
t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(0)
r = t.add_run("DemoTrack"); r.bold = True; r.font.size = Pt(38); r.font.color.rgb = RGBColor.from_string(PURPLE)
sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("Build Plan — Phases & Detailed Tasks"); r.italic = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(PURPLE_DK)
meta = doc.add_paragraph(); meta.paragraph_format.space_after = Pt(10); bottom_rule(meta, BORDER, "6")
r = meta.add_run("Companion to the Project Specification v3.1  ·  Tunis, Tunisia  ·  2026")
r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string(GREY)
body("This is the engineering checklist for building DemoTrack. It breaks every phase into concrete, "
     "checkable tasks, with what each phase depends on and what “done” means. Phases are built in order so "
     "the app is **usable early** — by the end of Phase 4 you can send and record a real demo — and grow in "
     "value from there. The Contact-Database research (Phase 0) runs in parallel throughout.", space_after=8)
callout("How to use this.",
        "Work top-to-bottom within each phase; the ☐ items are designed to be ticked off. A phase is shippable "
        "on its own, so deploy at the end of each. “Size” in the overview is a rough effort signal (S / M / L), "
        "not a deadline.", NOTE_BG, BLUE)

# overview table
h4("Phase overview")
ov = [
    ["0 · Contact DB Research", "Parallel", "A real, qualified label database", "ongoing"],
    ["1 · Foundation", "L", "Auth + schema + deploy", "—"],
    ["2 · Track Vault", "S", "Store your tracks", "—"],
    ["3 · Contacts CRM", "M", "Labels live in-app", "—"],
    ["4 · Send Demo", "M", "Send + record a demo", "★ first usable"],
    ["5 · Email Presets + AI Hook", "M", "1-minute tailored email", "—"],
    ["6 · Follow-up + Digest", "L", "Told who to chase, reminded", "—"],
    ["7 · Feedback Log", "S", "Every reply + silence captured", "—"],
    ["8 · A&R Intel + Prioritization", "M", "Ranked by history, intel attached", "—"],
    ["9 · Dashboard + Funnel", "M", "“What now?” + conversion", "—"],
    ["10 · Label Discovery", "M", "Find next target, add to CRM", "—"],
    ["11 · Link Tracking (optional)", "M", "See Form/DM opens", "—"],
    ["12 · Work Sessions + Goals", "S", "Accountability, derived", "—"],
    ["13 · Artist Press Kit", "M", "Auto-attaching EPK", "—"],
]
make_table(["Phase", "Size", "Outcome", "Milestone"], ov, widths=[2.4, 0.7, 3.0, 1.0], font=9)

page_break()

# ============================ PHASE 0 ============================
h_section("Phase 0 · Contact Database Research", "")
body("*Parallel track — runs throughout the whole build. The app is the engine; this database is the fuel.*")
meta_box("Build and maintain a curated, verified library of career-changing labels + promo contacts.",
         "Nothing — start immediately, alongside Phase 1.",
         "The labels table is populated and feeds Discovery & the CRM.")
h4("Set up")
tasks([
    "Create a research tracking sheet with columns matching the `labels` schema (name, tier, access_path, method, link, genre_tags, bpm range, 2 sources, last_verified, submission_requirements, parent_label_id, notes).",
    "Write the qualification rubric as a checklist: the **≥3-of-5 bar** (Beatport chart presence · booking impact · streaming credibility · scene respect · accessible to unknowns).",
])
h4("The research loop (per contact)")
tasks([
    "**Source** — mine Beatport charts, label rosters, artist credits, scene press, and submission directories; include non-label channels (radio, blogs, playlist/repost curators, premiere channels).",
    "**Verify** — find the real submission route (email / portal / DM) and confirm it against **2 independent sources**; discard dead inboxes and outdated forms.",
    "**Qualify** — apply the ≥3/5 bar; if it fails, it does not go in.",
    "**Tier + access_path + priority** — assign Elite / A / B and a realistic access path (cold-demo-friendly / open-window-only / needs-warm-intro / relationship-only).",
    "**Record** — capture submission requirements (format, unreleased-only, single-vs-album), genre/BPM fit, the contact link, the 2 sources, and an A&R-intel note.",
    "**Capture sub-label links** — note efficiency multipliers (1 Defected ≈ 11 sub-labels; 1 Repopulate Mars ≈ 3) via `parent_label_id`.",
    "**Maintain** — stamp `last_verified`; set a re-check cadence so Discovery can flag stale entries.",
])
h4("Migration")
tasks([
    "Export the sheet to CSV in the exact `labels` column order.",
    "Import ≈170 verified labels + ≈98 non-label promo contacts into the `labels` table (Supabase CSV import or a seed script).",
    "Re-confirm the 17 curated seed labels are present and correctly tiered.",
])
callout("Definition of done (ongoing).", "Never fully finished, but the in-app database is ready to drive "
        "Discovery and the CRM once migrated, and a maintenance cadence is in place.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 1 ============================
h_section("Phase 1 · Foundation", "")
meta_box("Stand up the project: data, auth, security, app shell, and deployment.",
         "Nothing (and Phase 0 in parallel).",
         "You can log in via magic link, your data is private by RLS, and the app is deployed.")
h4("Project setup")
tasks([
    "Create the Supabase project; record the project URL and anon/service keys.",
    "Scaffold the frontend: **Vite + React 19 + TypeScript**.",
    "Install and configure **Tailwind CSS 4**; set a base theme (brand purple, mobile-first defaults).",
    "Create `.env` for `SUPABASE_URL`, `SUPABASE_ANON_KEY`; keep `ANTHROPIC_API_KEY` **server-side only** (never in the client bundle).",
    "Add the Supabase JS client and a session/auth context.",
])
h4("Database schema (12 tables)")
tasks([
    "Write SQL migrations for all 12 tables: `contacts`, `tracks`, `submissions`, `feedback`, `templates`, `ar_intel`, `work_sessions`, `goals`, `link_events`, `press_kit`, `labels`, `notifications`.",
    "Define enums / controlled vocabularies: **genre_tags**, track **status**, submission **status** & **method**, **access_path**, **relationship_stage**, **mood**, **goal type/timeframe**, feedback **response_type**.",
    "Add foreign keys (e.g. `contacts.label_id → labels.id`, `submissions.track_id`/`contact_id`, `tracks.exclusive_hold_contact_id`).",
    "Add a `user_id` (owner) column to every user-owned table, defaulting to `auth.uid()`.",
])
h4("Security (RLS)")
tasks([
    "Enable Row-Level Security on every table.",
    "Write owner-only policies: `user_id = auth.uid()` for select/insert/update/delete on all personal tables.",
    "Make `labels` read-only to authenticated users (it's a shared library, no per-row owner).",
    "Create the **public press-kit RPC**: a `SECURITY DEFINER` function that takes a slug and returns only whitelisted public fields — the public page never reads the table directly.",
])
h4("App shell & auth")
tasks([
    "Configure **magic-link (email OTP)** auth in Supabase; build the login screen.",
    "Set up routing (React Router) with a protected-route guard and a main layout/nav.",
    "Add the **PWA**: web manifest, service worker, installable, mobile-first responsive base.",
])
h4("Deploy")
tasks([
    "Deploy to **Vercel**; set environment variables in the Vercel project.",
    "Connect the `demotrack.app` domain (and reserve the `press.` / `track.` subdomains).",
    "Verify: sign in via magic link end-to-end; confirm RLS blocks cross-account reads.",
])
callout("Definition of done.", "Log in via magic link, see an empty but secured account, on a deployed URL.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 2 ============================
h_section("Phase 2 · Track Vault", "")
meta_box("Your catalogue: store tracks, statuses, listen links, and holds.", "Phase 1.",
         "You can store your tracks and see their status.")
h4("Tasks")
tasks([
    "Build the track **list view**: status badges, sort, and filter by status/genre.",
    "Build the **add/edit form**: title, genre tags (multi-select from the vocab), BPM, key, status, notes.",
    "Add `listen_link` + `link_type` (`soundcloud-private` / `dropbox` / `other`); **audio is never uploaded** — only the link is stored.",
    "Implement the **readiness checklist** (mastered? · correct format? · metadata/tags? · artwork?) shown before a track becomes “demo-ready.”",
    "Add **exclusive-hold** fields (`exclusive_hold_contact_id`, `hold_until`) with a clear on/off display.",
    "Wire the status workflow `idea → demo-ready → submitted → signed` with valid transitions.",
    "Add empty states and form validation.",
])
callout("Definition of done.", "Add, edit, and list tracks with status; set a hold on a track.", DONE_BG, GREEN)

# ============================ PHASE 3 ============================
h_section("Phase 3 · Contacts CRM", "")
meta_box("Your people: every label/contact with submission details and history.", "Phase 1 (and Phase 0 data).",
         "Your label list lives in the app with submission details.")
h4("Tasks")
tasks([
    "Build the contacts **list** with the 7 categories (Labels, DJs, A&Rs, Curators, Blogs, Promoters, Radio) and category/access-path filters.",
    "Build the contact **form**: name, category, submission method, email/portal/DM link, `relationship_stage`, `access_path`, notes.",
    "**Seed the 17 curated labels** into `labels`, and provide a way to add them to your `contacts`.",
    "Build the contact **detail view** with a full history panel (submissions + feedback) — empty for now, wired to populate later.",
    "Display the derived **“times contacted”** count on each contact (count of submissions).",
])
callout("Definition of done.", "Your label list lives in the app with submission methods and a history view.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 4 ============================
h_section("Phase 4 · Send Demo", "")
body("*★ The first usable milestone — after this, the app pays you back.*")
meta_box("The submission flow: pick, prepare, send, record — under 3 minutes.", "Phases 2 & 3.",
         "You can send a demo and it's recorded.")
h4("Tasks")
tasks([
    "Build the flow: **select track → select contact → choose preset** (preset is a stub here; the engine arrives in Phase 5).",
    "Implement **method routing** with color-coded badges: Email 🟢 / Form 🟡 / DM 🔵.",
    "**Email** — build a `mailto:` / Gmail-compose deep link with prefilled subject + body; opens your own inbox.",
    "**Form** — open the portal/platform in a new tab.",
    "**DM** — open the profile in a new tab.",
    "Render the **pre-send checklist** from the contact/label `submission_requirements`.",
    "Add the **exclusive-hold warning** — block/warn if the chosen track is held by another label.",
    "On send, **record a `submissions` row** (status = `sent`, method, dates).",
    "Add the **confirm step** for Form/DM (“Did you send it?”) so it's logged accurately.",
    "Add the **auto-attach press-kit link** toggle (stubbed until Phase 13).",
])
callout("Definition of done.", "Send a demo through any of the 3 methods and have it recorded in history.", DONE_BG, GREEN)

# ============================ PHASE 5 ============================
h_section("Phase 5 · Email Presets + AI Hook", "")
meta_box("Templates that auto-fill, with AI suggesting only the one personal sentence.", "Phase 4.",
         "A tailored, personal demo email in ~1 minute, sent from your own inbox.")
h4("Presets & merge fields")
tasks([
    "Build **preset CRUD** (a template editor) storing skeletons with `{merge_fields}` and one `{hook}` slot.",
    "Build the **merge-field engine**: parse `{...}` and auto-fill from the track, contact, and your profile.",
    "Seed starter presets: **cold email**, **DM** (shorter), **warm follow-up**, **form-portal note**.",
])
h4("AI hook")
tasks([
    "Create a **serverless/Edge function** that calls the Anthropic API (key stays server-side).",
    "Write the **hook prompt** with guardrails: ≤30 words, **never fabricate** a release/stat/angle, use `ar_intel` only if present, otherwise return an honest generic line.",
    "Build the hook UI: generate → show 1–2 options → edit → **required** validation → **duplicate-send warning**.",
    "Add output handling + a plain-text fallback if the model's formatting varies.",
    "**Wire it into Send Demo**: the assembled subject + body now flow into the mailto/compose link.",
])
callout("Definition of done.", "Pick track → pick label → confirm one AI-suggested sentence → send, in about a minute.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 6 ============================
h_section("Phase 6 · Follow-up + Notifications Digest", "")
meta_box("Tell you who to chase, let you close the loop in one tap, and reach you when the app is shut.",
         "Phase 4 (submissions exist).",
         "The app tells you who to follow up and reminds you outside the app.")
h4("Status & queue")
tasks([
    "Implement the `submissions.status` state machine: `sent → opened → replied → considering → signed → passed`.",
    "Build the **follow-up queue** showing only `sent`/`opened` items with days-since.",
    "Add **one-tap actions** per row: **“Got a reply”** (opens a 5-second feedback log) and **“Still silent”** (keeps nudging).",
])
h4("Automation (pg_cron)")
tasks([
    "Schedule a job that queues a **7-day silence** nudge.",
    "Schedule a job that flags **14-day overdue** and **auto-logs a `no-response` feedback record**.",
    "Build follow-up preset templates (reuse the Phase-5 engine).",
])
h4("Digest")
tasks([
    "Create the `notifications` table + user prefs (channel, cadence).",
    "Build a **Scheduled Edge Function** that composes the daily/weekly digest (“3 follow-ups due, 1 demo seen”).",
    "Send the digest **to your own inbox** (Resend or Supabase SMTP) — note this is self-notification, so deliverability is a non-issue.",
])
callout("Definition of done.", "The overdue queue is accurate, one tap advances a send's status, and a scheduled digest lands in your inbox.", DONE_BG, GREEN)

# ============================ PHASE 7 ============================
h_section("Phase 7 · Feedback Log", "")
meta_box("Capture every reply — and every silence — as data.", "Phase 6.",
         "Every reply and non-reply is captured and searchable.")
h4("Tasks")
tasks([
    "Build **feedback CRUD** tied to a submission (and thus a track + contact).",
    "Wire the **one-tap reply logging** from the follow-up queue: response-type dropdown + optional pasted text.",
    "Surface the **auto `no-response` records** (created by Phase-6 cron) in the log.",
    "Define the **response-type taxonomy**: signed · pass · not-for-us · constructive · no-response · other.",
    "Add list + search and back-links to the track and contact.",
])
callout("Definition of done.", "Every reply and silence is captured, searchable, and linked back to its track and contact.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 8 ============================
h_section("Phase 8 · A&R Intel + Prioritization", "")
meta_box("Attach research to each contact and surface untried targets first — no scoring engine.",
         "Phases 3 & 5.", "Contacts show history at a glance, with intel attached and untried targets first.")
h4("Tasks")
tasks([
    "Build **A&R-intel CRUD** (one record per contact): key people, what they sign now, dos & don'ts, recent releases, your notes.",
    "**Feed `ar_intel` into the Phase-5 hook prompt** so the AI's sentence is informed.",
    "Display the **“times contacted”** count prominently on labels and contacts.",
    "Implement the **prioritization sort**: cold-demo-friendly + untried (0 contacts) first, then by tier.",
    "Confirm there is **no weighted score** — prioritization is access_path + times-contacted only.",
])
callout("Definition of done.", "Each contact shows its history and intel, and untried, reachable labels rise to the top.", DONE_BG, GREEN)

# ============================ PHASE 9 ============================
h_section("Phase 9 · Dashboard + Funnel", "")
meta_box("The home screen that answers “what should I do right now?” plus your real conversion.",
         "Phases 4, 6, 7.", "The home screen aggregates every signal and shows the funnel.")
h4("Tasks")
tasks([
    "Build the **dashboard** aggregating: sends due, follow-ups due today, recent opens/replies, goal progress, recent feedback, work-session streak.",
    "Build the **funnel** view (sent → opened → replied → considering → signed) backed by **Postgres views**.",
    "Add **response-rate by genre/tier**.",
    "Add **expected-vs-actual framing** (5–15% cold-reply prior) so a quiet month reads as “keep going.”",
    "Add **CSV export** of the core tables (your data, portable).",
])
callout("Definition of done.", "Opening the app immediately answers “what now?”, and the funnel shows honest conversion.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 10 ============================
h_section("Phase 10 · Label Discovery", "")
meta_box("Browse the curated library to find your next target and pull it into your CRM.",
         "Phases 3 & 0 (populated labels).", "You can browse/filter labels and add one to your CRM.")
h4("Tasks")
tasks([
    "Build the **browse UI** over the read-only `labels` master.",
    "Add **filters**: genre, BPM range, submission method, tier, access_path, last_verified.",
    "Add a **“Have I submitted?”** indicator (join to your submissions).",
    "Add **freshness/stale warnings** when `last_verified` is past the threshold.",
    "Build **“Add to my CRM”** — copy the label into `contacts` with a `label_id` back-reference.",
    "Flag **“source updated — review”** on your CRM copy when the master changes.",
    "Show **sub-label reach hints** via `parent_label_id` (“this 1 send reaches 11”).",
])
callout("Definition of done.", "Filter the library to your next target and add it to your CRM in one click.", DONE_BG, GREEN)

# ============================ PHASE 11 ============================
h_section("Phase 11 · Demo Link Tracking  (optional)", "")
body("*Lower priority — unlisted SoundCloud already gives a free listen signal. Build only if you want automated open data.*")
meta_box("See when a label opens a Form/DM demo link.", "Phase 4 (and the Edge runtime).",
         "You can see when a label opens your Form/DM demo.")
h4("Tasks")
tasks([
    "Create the `link_events` table.",
    "Build the **Edge Function redirect** at `track.demotrack.app/{hash}`.",
    "Generate a **unique hash per submission** — **Form/DM links only, never email**.",
    "Log each click: timestamp, repeat-open count, approximate country (IP geo).",
    "Feed **“seen, no reply”** triggers into the follow-up queue.",
    "Show open/repeat counts on the submission detail.",
])
callout("Definition of done.", "Opening a tracked Form/DM link is logged and can trigger a smart follow-up.", DONE_BG, GREEN)

page_break()

# ============================ PHASE 12 ============================
h_section("Phase 12 · Work Sessions + Goals", "")
meta_box("Gentle accountability — with progress derived, not typed.", "Phase 4 (submissions to count).",
         "You can log studio time and track progress against targets.")
h4("Work sessions")
tasks([
    "Build a **quick-add session**: hours + a **one-tap mood** (fire / good / ok / low).",
    "**Auto-derive “demos sent that day”** from submissions — don't ask the user.",
    "Compute and show a **session streak**.",
])
h4("Goals")
tasks([
    "Build goals: type (send / produce / response / growth / release / network), timeframe (monthly / quarterly / yearly), target.",
    "**Derive progress** with a Postgres view over `submissions` / `tracks` — the user only sets the target.",
    "Show progress bars on the dashboard.",
])
callout("Definition of done.", "Log a session in seconds and watch goal progress fill itself from real activity.", DONE_BG, GREEN)

# ============================ PHASE 13 ============================
h_section("Phase 13 · Artist Press Kit", "")
meta_box("A shareable public EPK that attaches itself to every demo.", "Phases 1 & 5 (AI) & Storage.",
         "You have a shareable EPK that auto-attaches to demos.")
h4("Tasks")
tasks([
    "Build **press_kit CRUD**: artist name, location, genres, bio, past releases, links (SoundCloud/Beatport/Instagram), follower/play counts, photo, slug, `is_public`.",
    "Add **photo upload** to Supabase Storage (the only hosted asset).",
    "Wire **AI bio generation** in 3 tones (facts-only, no invented stats).",
    "Keep stats **manual** with an **“updated {date}”** stamp.",
    "Build the **public page** at `press.demotrack.app/{slug}`, served by the **whitelisted slug RPC** (no direct table reads).",
    "Wire the **auto-attach toggle** into Send Demo (Phase 4).",
])
callout("Definition of done.", "A clean public EPK exists at your slug and attaches its link to every demo by default.", DONE_BG, GREEN)

page_break()

# ============================ APPENDIX ============================
h_section("Appendix · Build sequencing notes", "")
h4("Why this order")
bullets([
    "**Usable early:** Phases 1–4 get you to a working send-and-record tool; everything after makes each send smarter.",
    "**Each phase ships:** deploy at the end of every phase — no big-bang release.",
    "**Parallelism:** Phase 0 (research) runs the entire time; the app is useless without the database, and the database takes real time to build.",
])
h4("Hard dependencies at a glance")
make_table(["Phase", "Cannot start until"],
    [["4 Send Demo", "2 Track Vault + 3 Contacts CRM"],
     ["5 AI Hook", "4 Send Demo"],
     ["6 Follow-up", "4 Send Demo (submissions exist)"],
     ["7 Feedback", "6 Follow-up (one-tap logging)"],
     ["8 Intel + Prioritization", "3 Contacts + 5 (to feed the hook)"],
     ["9 Dashboard", "4 + 6 + 7 (signals to aggregate)"],
     ["10 Discovery", "0 (populated labels) + 3"],
     ["13 Press Kit", "1 + 5 (AI bio) + Storage"]],
    widths=[2.4, 4.7], font=9.5)
h4("Carry these principles into every phase")
bullets([
    "**Derive, don't ask** — compute counts and progress; never make the user type a number.",
    "**One-tap over typing** — free text only for the hook, A&R notes, and feedback.",
    "**Silence is data** — overdue sends log themselves.",
    "**Structural privacy** — public data only ever leaves through the whitelisted slug endpoint.",
])

foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(16)
r = foot.add_run("DemoTrack · Build Plan · Companion to Project Specification v3.1 · Tunis, 2026")
r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string("9AA0AB")

out = "/home/user/DemoTrack/docs/DemoTrack_Build_Plan.docx"
doc.save(out)
print("Saved", out)
